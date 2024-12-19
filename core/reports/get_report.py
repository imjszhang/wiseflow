import random
import re
import os
import json
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from llms.openai_wrapper import openai_llm
from llms.dify_wrapper import dify_llm
from loguru import logger
from utils.general_utils import get_logger_level, isChinesePunctuation
from utils.pb_api import PbTalker
from utils.prompt_utils import load_prompt_template
import uuid


class ReportService:
    """
    报告生成服务类，用于根据 insight 和相关文章生成报告并保存为 Word 文档。
    """

    def __init__(self):
        """
        初始化 ReportService，包括日志、缓存目录和内存存储。
        """
        self.project_dir = os.environ.get("PROJECT_DIR", "")
        # 初始化缓存目录
        self.cache_url = os.path.join(self.project_dir, 'backend_service')
        os.makedirs(self.cache_url, exist_ok=True)

        # 内存存储，用于缓存 insight 的历史数据
        self.memory = {}

        logger.info('Backend service initialized successfully.')

    async def report(self, insight_id: str, topics: list[str], comment: str) -> dict:
        """
        根据 insight ID 和相关文章生成报告。

        :param insight_id: str - insight 的唯一标识符
        :param topics: list[str] - 报告的主题列表
        :param comment: str - 用户的附加评论
        :return: dict - 包含生成状态和结果的字典
        """
        logger.debug(f'Got new report request for insight_id {insight_id}')
        insight = pb.read('insights', filter=f'id="{insight_id}"')
        if not insight:
            logger.error(f'Insight {insight_id} not found')
            return self.build_out(-2, 'Insight not found')

        article_ids = insight[0]['articles']
        if not article_ids:
            logger.error(f'Insight {insight_id} has no articles')
            return self.build_out(-2, 'Cannot find articles for insight')

        # 获取相关文章
        article_list = [pb.read('articles', fields=['title', 'abstract', 'content', 'url', 'publish_time'], filter=f'id="{_id}"')
                        for _id in article_ids]
        article_list = [_article[0] for _article in article_list if _article]

        if not article_list:
            logger.debug(f'{insight_id} has no valid articles')
            return self.build_out(-2, f'{insight_id} has no valid articles')

        content = insight[0]['content']
        memory = self.memory.get(insight_id, '')

        # 生成报告并保存为 Word 文档
        docx_file = os.path.join(self.cache_url, f'{insight_id}_{uuid.uuid4()}.docx')
        flag, memory = await get_report(content, article_list, memory, topics, comment, docx_file)
        self.memory[insight_id] = memory

        if flag:
            # 上传生成的报告到 PB
            file = open(docx_file, 'rb')
            message = pb.upload('insights', insight_id, 'docx', f'{insight_id}.docx', file)
            file.close()
            if message:
                logger.debug(f'Report successfully generated and uploaded: {message}')
                return self.build_out(11, message)
            else:
                logger.error(f'{insight_id} report generated successfully, but failed to upload to PB.')
                return self.build_out(-2, 'Report generated successfully, but failed to upload to PB.')
        else:
            logger.error(f'{insight_id} failed to generate report.')
            return self.build_out(-11, 'Report generation failed.')

    def build_out(self, flag: int, answer: str = "") -> dict:
        """
        构建返回结果。

        :param flag: int - 状态码
        :param answer: str - 返回的消息
        :return: dict - 包含状态码和消息的字典
        """
        return {"flag": flag, "result": [{"type": "text", "answer": answer}]}


# 初始化日志
project_dir = os.environ.get("PROJECT_DIR", "")
if project_dir:
    os.makedirs(project_dir, exist_ok=True)
logger_file = os.path.join(project_dir, 'wiseflow.log')
dsw_log = get_logger_level()
logger.add(
    logger_file,
    level=dsw_log,
    backtrace=True,
    diagnose=True,
    rotation="50 MB"
)

# 初始化 PbTalker
pb = PbTalker(logger)

# 最大输入 token 限制
max_input_tokens = 30000

# 角色设定
character = '来自中国的网络安全情报专家'
report_type = '网络安全情报'

# 动态选择 LLM 提供商
LLM_PROVIDER = os.environ.get("LLM_PROVIDER", "dify").lower()

def select_llm():
    """
    根据环境变量选择 LLM 提供商。
    """
    if LLM_PROVIDER == "openai":
        return openai_llm
    elif LLM_PROVIDER == "dify":
        return dify_llm
    else:
        raise ValueError(f"Unsupported LLM provider: {LLM_PROVIDER}")

# 动态选择 LLM
llm = select_llm()

# 定义提示词模板文件路径
PROMPT_DIR = "prompts"
PROMPT_DIR = os.path.join(PROMPT_DIR, "reports")
language = "zh"  # 假设语言为中文
PROMPT_DIR = os.path.join(PROMPT_DIR, language)

SYSTEM_PROMPT_FILE = os.path.join(PROMPT_DIR, "system_prompt.txt")
USER_PROMPT_FILE = os.path.join(PROMPT_DIR, "user_prompt.txt")

# 加载提示词模板
try:
    system_prompt_template = load_prompt_template(SYSTEM_PROMPT_FILE)
    user_prompt_template = load_prompt_template(USER_PROMPT_FILE)
except FileNotFoundError as e:
    logger.error(f"Failed to load prompt templates: {e}")
    exit(1)

async def get_report(insight: str, articles: list[dict], memory: str, topics: list[str], comment: str, docx_file: str):
    """
    根据 insight 和相关文章生成报告并保存为 Word 文档。

    :param insight: str - insight 的内容
    :param articles: list[dict] - 相关文章列表
    :param memory: str - 历史报告内容
    :param topics: list[str] - 报告的主题列表
    :param comment: str - 用户的附加评论
    :param docx_file: str - 保存的 Word 文档路径
    :return: tuple[bool, str] - 是否成功生成报告和生成的报告内容
    """
    zh_index = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十', '十一', '十二']

    if isChinesePunctuation(insight[-1]):
        insight = insight[:-1]

    # 分离段落和标题
    if len(topics) == 0:
        title = ''
    elif len(topics) == 1:
        title = topics[0]
        topics = []
    else:
        title = topics[0]
        topics = [s.strip() for s in topics[1:] if s.strip()]

    schema = f'【标题】{title}\n\n【综述】\n\n'
    if topics:
        for i in range(len(topics)):
            schema += f'【{zh_index[i]}、{topics[i]}】\n\n'

    # 判断是否是修改要求
    system_prompt, user_prompt = '', ''
    if memory and comment:
        paragraphs = re.findall("、(.*?)】", memory)
        if set(topics) <= set(paragraphs):
            logger.debug("No change in Topics, need to modify the report")
            system_prompt = system_prompt_template.format(
                character=character,
                report_type=report_type,
                memory=memory
            )
            user_prompt = user_prompt_template.format(comment=comment)

    if not system_prompt or not user_prompt:
        logger.debug("Need to generate the report")
        texts = ''
        for article in articles:
            if article['content']:
                texts += f"<article>{article['content']}</article>\n"
            else:
                if article['abstract']:
                    texts += f"<article>{article['abstract']}</article>\n"
                else:
                    texts += f"<article>{article['title']}</article>\n"

            if len(texts) > max_input_tokens:
                break

        logger.debug(f"Articles context length: {len(texts)}")
        system_prompt = system_prompt_template.format(
            character=character,
            report_type=report_type,
            insight=insight,
            articles=texts
        )

        if comment:
            user_prompt = user_prompt_template.format(comment=comment, schema=schema)
        else:
            user_prompt = user_prompt_template.format(comment="", schema=schema)

    # 调用大模型生成报告
    check_flag = False
    check_list = schema.split('\n\n')
    check_list = [_[1:] for _ in check_list if _.startswith('【')]
    result = ''
    for i in range(2):
        try:
            if LLM_PROVIDER == "openai":
                result = llm(
                    [{'role': 'system', 'content': system_prompt}, {'role': 'user', 'content': user_prompt}],
                    model="gpt-4o-mini-2024-07-18",
                    temperature=0.1,
                    logger=logger
                )
            elif LLM_PROVIDER == "dify":
                inputs = {'system': system_prompt}
                response = llm(user_prompt, 'wiseflow', inputs=inputs, logger=logger)
                result = response['answer']
            else:
                raise ValueError(f"Unsupported LLM provider: {LLM_PROVIDER}")
        except Exception as e:
            logger.error(f"Error during LLM call: {e}")
            continue

        logger.debug(f"Raw result:\n{result}")
        if len(result) > 50:
            check_flag = True
            for check_item in check_list[2:]:
                if check_item not in result:
                    check_flag = False
                    break

        if check_flag:
            break

        logger.debug("Result not good, re-generating...")

    if not check_flag:
        if not result:
            logger.warning('Report-process-error: LLM out of work!')
            return False, ''
        else:
            logger.warning('Report-process-error: Cannot generate, change topics and insight, then re-try')
            return False, ''

    # 解析生成的报告
    contents = result.split("【")
    bodies = {}
    for text in contents:
        for item in check_list:
            if text.startswith(item):
                check_list.remove(item)
                key, value = text.split("】")
                value = value.strip()
                if isChinesePunctuation(value[0]):
                    value = value[1:]
                bodies[key] = value.strip()
                break

    if not bodies:
        logger.warning('Report-process-error: Cannot generate, change topics and insight, then re-try')
        return False, ''

    if '标题' not in bodies:
        if "】" in contents[0]:
            _title = contents[0].split("】")[0]
            bodies['标题'] = _title.strip()
        else:
            if len(contents) > 1 and "】" in contents[1]:
                _title = contents[0].split("】")[0]
                bodies['标题'] = _title.strip()
            else:
                bodies['标题'] = ""

    # 生成 Word 文档
    doc = Document()
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = Pt(12)
    doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

    # 添加标题和摘要
    if not title:
        title = bodies['标题']

    Head = doc.add_heading(level=1)
    Head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = Head.add_run(title)
    run.font.name = u'Cambria'
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')

    doc.add_paragraph(
        f"\n生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    del bodies['标题']
    if '综述' in bodies:
        doc.add_paragraph(f"\t{bodies['综述']}\n")
        del bodies['综述']

    # 添加章节内容
    for key, value in bodies.items():
        Head = doc.add_heading(level=2)
        run = Head.add_run(key)
        run.font.name = u'Cambria'
        run.font.color.rgb = RGBColor(0, 0, 0)
        doc.add_paragraph(f"{value}\n")

    # 添加附件引用信息源
    Head = doc.add_heading(level=2)
    run = Head.add_run("附：原始信息网页")
    run.font.name = u'Cambria'
    run.font.color.rgb = RGBColor(0, 0, 0)

    contents = []
    for i, article in enumerate(articles):
        date_text = str(article['publish_time'])
        if len(date_text) == 8:
            date_text = f"{date_text[:4]}-{date_text[4:6]}-{date_text[6:]}"

        contents.append(f"{i+1}、{article['title']}|{date_text}\n{article['url']} ")

    doc.add_paragraph("\n\n".join(contents))

    doc.save(docx_file)

    return True, result[result.find("【"):]